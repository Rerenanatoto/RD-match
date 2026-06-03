import io
import re
import math
import unicodedata
from datetime import date
from urllib.parse import urlparse

import pandas as pd
import streamlit as st
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.text.paragraph import Paragraph
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment
from openpyxl.utils import get_column_letter
from rapidfuzz import fuzz, process

st.set_page_config(page_title="RD Pipeline + DDQ", layout="wide")
st.title("RD Pipeline + DDQ – Extração · Consolidação · Fontes · Word")

# ── regex ─────────────────────────────────────────────────────────────────────
URL_RE = re.compile(r"https?://[^\s\)\]\}\"\'>]+", re.IGNORECASE)

# ── DDQ constants ─────────────────────────────────────────────────────────────
DDQ_COR_AZUL       = RGBColor(31, 78, 121)
DDQ_COR_AZUL_CLARO = RGBColor(79, 129, 189)
DDQ_COR_CINZA      = RGBColor(90, 90, 90)
DDQ_COR_PRETO      = RGBColor(0, 0, 0)
DDQ_COR_LINHA      = "D9E2F3"
DDQ_COR_VERMELHO   = RGBColor(255, 0, 0)
DDQ_REFERENCIA_DATA_DEFAULT = "February 2026"

# =============================================================================
# ── helpers gerais (RD Pipeline) ─────────────────────────────────────────────
# =============================================================================

def extract_urls(text: str):
    return URL_RE.findall(text or "")


def is_source_line(text: str) -> bool:
    t = (text or "").strip().lower()
    if not t:
        return False
    if t.startswith("same from above"):
        return True
    if t.startswith("source:"):
        return True
    return bool(extract_urls(t))


def append_sources(data, idx: int, sources):
    if idx is None or not sources:
        return
    seen, ordered = set(), []
    for s in sources:
        s = (s or "").strip()
        if not s:
            continue
        if s.lower().startswith("source:"):
            s = s.split(":", 1)[1].strip()
        urls = extract_urls(s)
        if urls:
            for u in urls:
                if u not in seen:
                    seen.add(u); ordered.append(u)
        else:
            if s not in seen:
                seen.add(s); ordered.append(s)
    if not ordered:
        return
    existing = (data[idx][1] or "").strip()
    if existing:
        ex_lines = [ln.strip() for ln in existing.split("\n") if ln.strip()]
        for ln in ex_lines:
            if ln not in seen:
                seen.add(ln)
        data[idx][1] = existing + "\n" + "\n".join([s for s in ordered if s not in ex_lines])
    else:
        data[idx][1] = "\n".join(ordered)


def ajustar_altura(ws, cols=(1, 2), largura_coluna=100, altura_por_linha=15):
    for col in cols:
        ws.column_dimensions[get_column_letter(col)].width = largura_coluna
    for row in range(1, ws.max_row + 1):
        max_lines = 1
        for col in cols:
            cell = ws.cell(row=row, column=col)
            cell.alignment = Alignment(wrap_text=True, vertical="top", horizontal="left")
            if cell.value:
                texto = str(cell.value)
                lines = 0
                for part in texto.split("\n"):
                    if len(part) == 0:
                        lines += 1
                    else:
                        lines += max(1, math.ceil(len(part) / (largura_coluna * 1.2)))
                max_lines = max(max_lines, lines)
        ws.row_dimensions[row].height = max_lines * altura_por_linha


def _run_is_red(run) -> bool:
    try:
        c = run.font.color
        if c and c.rgb:
            return str(c.rgb).upper() == "FF0000"
    except Exception:
        pass
    return False


def is_red_paragraph(paragraph) -> bool:
    runs = [r for r in paragraph.runs if r.text.strip()]
    return bool(runs) and all(_run_is_red(r) for r in runs)


def clean_text_for_match(text: str) -> str:
    cleaned = URL_RE.sub("", text)
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    return cleaned.lower()


def insert_paragraph_after(paragraph, text=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para


# =============================================================================
# ── Etapa 1: docx → xlsx (RD Pipeline) ───────────────────────────────────────
# =============================================================================

def word_to_excel_bytes(docx_bytes: bytes) -> bytes:
    doc = Document(io.BytesIO(docx_bytes))
    data = []
    current_idx = None

    for paragraph in doc.paragraphs:
        raw  = paragraph.text
        text = (raw or "").strip()
        if not text:
            continue
        if paragraph.style and paragraph.style.name.startswith("Heading"):
            continue
        if paragraph._element.xpath("ancestor::w:tbl"):
            continue

        # ── verifica fonte ANTES dos filtros de formatação ──────────────────────
        # Detecta como fonte: (a) texto com padrão "Source:" / URL ou
        # (b) parágrafo em vermelho (inclui continuações sem prefixo "Source:")
        if is_source_line(text) or is_red_paragraph(paragraph):
            if current_idx is None:
                data.append(["", ""])
                current_idx = len(data) - 1
            append_sources(data, current_idx, [text])
            continue

        all_bold       = all(r.bold           for r in paragraph.runs if r.text.strip())
        all_italic     = all(r.italic         for r in paragraph.runs if r.text.strip())
        all_underlined = all(r.font.underline  for r in paragraph.runs if r.text.strip())
        if all_bold or all_italic or all_underlined:
            continue
        if text.startswith("\u201c") and text.endswith("\u201d"):
            continue

        if text.startswith("Note:") or re.match(r"^\(\d+\)", text):
            continue

        urls    = extract_urls(text)
        cleaned = URL_RE.sub("", text)
        cleaned = re.sub(r"\s+", " ", cleaned).strip()

        if cleaned:
            data.append([cleaned, ""])
            current_idx = len(data) - 1
            if urls:
                append_sources(data, current_idx, urls)
        else:
            if urls:
                if current_idx is None:
                    data.append(["", ""])
                    current_idx = len(data) - 1
                append_sources(data, current_idx, urls)

    wb = Workbook()
    ws = wb.active
    ws.title = "Parágrafos"
    for i, (t, s) in enumerate(data, start=1):
        ws.cell(row=i, column=1, value=t)
        ws.cell(row=i, column=2, value=s)
    ajustar_altura(ws, cols=(1, 2), largura_coluna=100, altura_por_linha=20)
    align = Alignment(wrap_text=True, vertical="top", horizontal="left")
    for row in ws.iter_rows(min_row=1, max_row=ws.max_row, min_col=1, max_col=2):
        for cell in row:
            cell.alignment = align

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.read()


# =============================================================================
# ── Etapa 2: combinar excels (RD Pipeline) ───────────────────────────────────
# =============================================================================

def combinar_excels(bytes_novo: bytes, bytes_antigo: bytes) -> bytes:
    df_novo   = pd.read_excel(io.BytesIO(bytes_novo),   header=None, engine="openpyxl")
    df_antigo = pd.read_excel(io.BytesIO(bytes_antigo), header=None, engine="openpyxl")
    df_final  = pd.concat([df_antigo, df_novo], ignore_index=True)
    buf = io.BytesIO()
    df_final.to_excel(buf, index=False, header=False)
    buf.seek(0)
    return buf.read()


# =============================================================================
# ── Etapa 3: aplicar fontes / fuzzy match (RD Pipeline) ──────────────────────
# =============================================================================

def apply_sources_bytes(docx_bytes: bytes, excel_bytes: bytes, threshold: int = 80):
    """Aplica fontes no Word usando fuzzy matching (rapidfuzz partial_ratio)."""
    wb = load_workbook(io.BytesIO(excel_bytes))
    ws = wb.active

    pares = []
    for row in ws.iter_rows(min_row=1, max_col=2, values_only=True):
        a, b = row[0], row[1]
        if not a or not b:
            continue
        chave = clean_text_for_match(str(a).strip())
        fonte = str(b).strip()
        if chave and fonte:
            pares.append((chave, fonte))

    doc = Document(io.BytesIO(docx_bytes))

    if not pares:
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read(), pd.DataFrame(columns=["Parágrafo", "Fonte aplicada", "Score"])

    chaves     = [p[0] for p in pares]
    fontes_map = {p[0]: p[1] for p in pares}
    preview    = []

    for p in doc.paragraphs:
        txt = p.text.strip()
        if not txt or is_red_paragraph(p):
            continue
        texto_par = clean_text_for_match(txt)
        resultado = process.extractOne(texto_par, chaves, scorer=fuzz.partial_ratio)
        score     = resultado[1] if resultado else 0

        if resultado and score >= threshold:
            fonte_val = fontes_map[resultado[0]]
            linhas    = [ln.strip() for ln in fonte_val.split("\n") if ln.strip()]
            last_p    = p
            for linha in linhas:
                new_p = insert_paragraph_after(last_p)
                run   = new_p.add_run(linha)
                run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                run.font.size      = Pt(10)
                new_p.paragraph_format.space_before = Pt(12)
                last_p = new_p
            preview.append({"Parágrafo": txt, "Fonte aplicada": fonte_val, "Score": score})
        else:
            preview.append({"Parágrafo": txt, "Fonte aplicada": "", "Score": score})

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read(), pd.DataFrame(preview)


# =============================================================================
# ── Etapa 4: excel → docx (RD Pipeline) ─────────────────────────────────────
# =============================================================================

def excel_para_docx_bytes(excel_bytes: bytes) -> bytes:
    df  = pd.read_excel(io.BytesIO(excel_bytes), header=None, engine="openpyxl")
    doc = Document()

    title = doc.add_paragraph("", style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Recent Developments – With Sources")

    for i in range(df.shape[0]):
        texto = str(df.iat[i, 0]).strip()
        fonte = str(df.iat[i, 1]).strip() if df.shape[1] > 1 else ""
        if not texto or texto.lower() in ("nan", ""):
            continue
        p_txt = doc.add_paragraph(texto)
        p_txt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_txt.paragraph_format.space_after = Pt(3)
        if fonte and fonte.lower() not in ("nan", ""):
            p_src = doc.add_paragraph()
            r     = p_src.add_run(fonte)
            r.font.size      = Pt(10)
            r.font.color.rgb = RGBColor(255, 0, 0)
            p_src.paragraph_format.space_after = Pt(12)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# =============================================================================
# ── helpers DDQ (texto / parse) ──────────────────────────────────────────────
# =============================================================================

def ddq_normalizar_texto(texto):
    if texto is None:
        return ""
    texto = str(texto).strip().lower()
    texto = unicodedata.normalize("NFKD", texto)
    texto = "".join(ch for ch in texto if not unicodedata.combining(ch))
    texto = re.sub(r"[^a-z0-9]+", "", texto)
    return texto


def ddq_limpar_texto(x):
    if pd.isna(x):
        return ""
    return str(x).strip()


def ddq_valor_incluir(x):
    if pd.isna(x):
        return False
    try:
        return float(x) == 1.0
    except Exception:
        pass
    txt = str(x).strip().lower()
    return txt in {"1", "1.0", "sim", "yes", "true"}


def ddq_localizar_colunas(df):
    mapa = {ddq_normalizar_texto(col): col for col in df.columns}
    candidatos = {
        "number":   ["number", "numero", "n"],
        "question": ["question", "pergunta"],
        "answer":   ["answer", "resposta"],
        "source":   ["source", "fonte"],
        "include":  [
            "include1ornot0",
            "include10ornot0",
            "include",
            "include1ornot0?"
        ],
    }
    colunas = {}
    for chave, opcoes in candidatos.items():
        encontrada = None
        for opcao in opcoes:
            if opcao in mapa:
                encontrada = mapa[opcao]
                break
        if not encontrada:
            for col_norm, col_real in mapa.items():
                if chave == "include":
                    if "include" in col_norm and "1" in col_norm and "0" in col_norm:
                        encontrada = col_real
                        break
                else:
                    if any(opcao in col_norm for opcao in opcoes):
                        encontrada = col_real
                        break
        colunas[chave] = encontrada

    faltantes = [k for k, v in colunas.items() if v is None]
    if faltantes:
        raise ValueError(
            f"Não foi possível localizar as colunas obrigatórias: {faltantes}. "
            f"Colunas encontradas: {list(df.columns)}"
        )
    return colunas


def ddq_normalizar_numero_pergunta(x):
    if pd.isna(x):
        return ""
    texto = str(x).strip()
    try:
        valor = float(texto)
        if valor.is_integer():
            return str(int(valor))
        return texto
    except Exception:
        return texto


def ddq_limpar_fonte_bruta(texto):
    if not texto:
        return ""
    texto = str(texto)
    texto = texto.replace("\\_", "_")
    texto = texto.replace("\\xa0", " ")
    texto = re.sub(r"\s+", " ", texto).strip()
    return texto


def ddq_extrair_markdown_links(texto):
    if not texto:
        return []
    padrao = re.compile(r"\[([^\]]+)\]\((https?://[^)\s]+)\)")
    encontrados = padrao.findall(texto)
    resultado = []
    vistos = set()
    for label, url in encontrados:
        url = url.strip().rstrip(".,;")
        chave = (label.strip(), url)
        if chave not in vistos:
            resultado.append((label.strip(), url))
            vistos.add(chave)
    return resultado


def ddq_extrair_urls_simples(texto):
    if not texto:
        return []
    urls = re.findall(r'https?://[^\s\]\)>,"\\]+', texto)
    resultado = []
    vistos = set()
    for url in urls:
        url = url.strip().rstrip(".,;")
        if url not in vistos:
            resultado.append(url)
            vistos.add(url)
    return resultado


def ddq_remover_links_do_texto(texto):
    if not texto:
        return ""
    texto = re.sub(r"\[([^\]]+)\]\((https?://[^)\s]+)\)", r"\1", texto)
    texto = re.sub(r'https?://[^\s\]\)>,"\\]+', "", texto)
    texto = texto.replace("[]", " ")
    texto = re.sub(r"\s+", " ", texto).strip(" -;,.[()]")
    return texto.strip()


def ddq_rotulo_amigavel_para_url(url):
    try:
        parsed = urlparse(url)
        dominio = parsed.netloc.lower().replace("www.", "")
        caminho = parsed.path.strip("/")
        mapa_dominios = {
            "bcb.gov.br":                   "BCB",
            "gov.br":                       "gov.br",
            "planalto.gov.br":              "Planalto",
            "tesourotransparente.gov.br":   "Tesouro Transparente",
            "thot-arquivos.tesouro.gov.br": "Tesouro",
            "agenciagov.ebc.com.br":        "Agência Gov",
        }
        rotulo_base = None
        for dom, nome in mapa_dominios.items():
            if dominio.endswith(dom):
                rotulo_base = nome
                break
        if rotulo_base is None:
            rotulo_base = dominio
        if caminho:
            ultimo = caminho.split("/")[-1]
            ultimo = ultimo.replace("-", " ").replace("_", " ").strip()
            if ultimo:
                return f"{rotulo_base} — {ultimo[:60]}"
        return rotulo_base
    except Exception:
        return url


def ddq_consolidar_fontes(texto_fonte):
    texto = ddq_limpar_fonte_bruta(texto_fonte)
    md_links = ddq_extrair_markdown_links(texto)
    urls_simples = ddq_extrair_urls_simples(texto)

    links = []
    urls_ja_usadas = set()

    for label, url in md_links:
        links.append({"label": label, "url": url})
        urls_ja_usadas.add(url)

    for url in urls_simples:
        if url not in urls_ja_usadas:
            links.append({"label": ddq_rotulo_amigavel_para_url(url), "url": url})
            urls_ja_usadas.add(url)

    texto_extra = ddq_remover_links_do_texto(texto)
    if texto_extra in {"OK", "OK[]", "[]", ""}:
        texto_extra = ""

    return {"links": links, "texto_extra": texto_extra}


# =============================================================================
# ── helpers DDQ (Word / formatação) ──────────────────────────────────────────
# =============================================================================

def ddq_set_paragraph_border_bottom(paragraph, color="D9D9D9", size="8", space="2"):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    bottom = pBdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        pBdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)


def ddq_shade_paragraph(paragraph, fill="F4F8FB"):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    shd = pPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        pPr.append(shd)
    shd.set(qn("w:fill"), fill)


def ddq_add_hyperlink(paragraph, text, url, color="0563C1", underline=True):
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    if color:
        c = OxmlElement("w:color")
        c.set(qn("w:val"), color)
        rPr.append(c)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)
    new_run.append(rPr)
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def ddq_configurar_documento(doc):
    sec = doc.sections[0]
    sec.top_margin    = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.5)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    h1 = styles["Heading 1"]
    h1.font.name      = "Calibri"
    h1.font.size      = Pt(15)
    h1.font.bold      = True
    h1.font.color.rgb = DDQ_COR_AZUL

    h2 = styles["Heading 2"]
    h2.font.name      = "Calibri"
    h2.font.size      = Pt(12)
    h2.font.bold      = True
    h2.font.color.rgb = DDQ_COR_AZUL_CLARO

    if "DDQ Label" not in styles:
        st_label = styles.add_style("DDQ Label", WD_STYLE_TYPE.PARAGRAPH)
        st_label.base_style     = styles["Normal"]
        st_label.font.name      = "Calibri"
        st_label.font.size      = Pt(10.5)
        st_label.font.bold      = True
        st_label.font.color.rgb = DDQ_COR_AZUL

    if "DDQ Corpo" not in styles:
        st_corpo = styles.add_style("DDQ Corpo", WD_STYLE_TYPE.PARAGRAPH)
        st_corpo.base_style     = styles["Normal"]
        st_corpo.font.name      = "Calibri"
        st_corpo.font.size      = Pt(11)
        st_corpo.font.color.rgb = DDQ_COR_PRETO

    if "DDQ Fonte" not in styles:
        st_fonte = styles.add_style("DDQ Fonte", WD_STYLE_TYPE.PARAGRAPH)
        st_fonte.base_style     = styles["Normal"]
        st_fonte.font.name      = "Calibri"
        st_fonte.font.size      = Pt(10)
        st_fonte.font.color.rgb = DDQ_COR_CINZA


def ddq_adicionar_texto_vermelho(doc, texto, keep_with_next=False):
    p = doc.add_paragraph(style="DDQ Corpo")
    p.paragraph_format.space_before   = Pt(6)
    p.paragraph_format.space_after    = Pt(6)
    p.paragraph_format.keep_with_next = keep_with_next
    run = p.add_run(texto)
    run.font.color.rgb = DDQ_COR_VERMELHO
    run.font.italic    = True


def ddq_adicionar_bloco_label_valor(doc, label, valor, sombrear_label=False):
    p_label = None
    if label:
        p_label = doc.add_paragraph(style="DDQ Label")
        p_label.paragraph_format.space_before = Pt(4)
        p_label.paragraph_format.space_after  = Pt(1)
        p_label.add_run(label)
        if sombrear_label:
            ddq_shade_paragraph(p_label, fill="F3F7FB")

    p_valor = doc.add_paragraph(style="DDQ Corpo")
    p_valor.paragraph_format.left_indent  = Cm(0.35)
    p_valor.paragraph_format.space_before = Pt(0)
    p_valor.paragraph_format.space_after  = Pt(5)
    p_valor.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p_valor.add_run(valor if valor else "")
    return p_label, p_valor


def ddq_adicionar_fontes_lapidadas(doc, texto_fonte):
    p_label = doc.add_paragraph(style="DDQ Label")
    p_label.paragraph_format.space_before = Pt(4)
    p_label.paragraph_format.space_after  = Pt(1)
    p_label.add_run("Source")

    dados       = ddq_consolidar_fontes(texto_fonte)
    links       = dados["links"]
    texto_extra = dados["texto_extra"]

    if links:
        for item in links:
            p = doc.add_paragraph(style="DDQ Fonte")
            p.paragraph_format.left_indent  = Cm(0.6)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(1)
            bullet = p.add_run("• ")
            bullet.font.color.rgb = DDQ_COR_CINZA
            ddq_add_hyperlink(p, item["label"], item["url"], color="0563C1", underline=True)
            if item["label"] != item["url"]:
                r = p.add_run(f" ({item['url']})")
                r.font.size      = Pt(9)
                r.font.color.rgb = DDQ_COR_CINZA

    if texto_extra:
        p_txt = doc.add_paragraph(style="DDQ Fonte")
        p_txt.paragraph_format.left_indent  = Cm(0.35)
        p_txt.paragraph_format.space_before = Pt(1)
        p_txt.paragraph_format.space_after  = Pt(4)
        p_txt.add_run(texto_extra)

    if not links and not texto_extra:
        p_txt = doc.add_paragraph(style="DDQ Fonte")
        p_txt.paragraph_format.left_indent  = Cm(0.35)
        p_txt.paragraph_format.space_before = Pt(0)
        p_txt.paragraph_format.space_after  = Pt(4)
        p_txt.add_run("—")

    return p_label


def ddq_adicionar_pergunta(doc, numero_real, pergunta, resposta, fallback_numero=None):
    numero_titulo = (
        numero_real if numero_real
        else str(fallback_numero) if fallback_numero is not None
        else ""
    )
    p_item = doc.add_paragraph(style="Heading 2")
    p_item.paragraph_format.space_before = Pt(8)
    p_item.paragraph_format.space_after  = Pt(4)
    p_item.add_run(f"Question {numero_titulo}" if numero_titulo else "Question")

    ddq_adicionar_bloco_label_valor(doc, "", resposta, sombrear_label=False)

    p_sep = doc.add_paragraph()
    p_sep.paragraph_format.space_before = Pt(6)
    p_sep.paragraph_format.space_after  = Pt(8)
    ddq_set_paragraph_border_bottom(p_sep, color=DDQ_COR_LINHA, size="10", space="2")


def ddq_montar_lista_perguntas_respondidas(df_filtrado, col_number):
    numeros = []
    for _, row in df_filtrado.iterrows():
        numero = ddq_normalizar_numero_pergunta(row.get(col_number))
        if numero:
            numeros.append(numero)
    return numeros


def ddq_adicionar_resumo_perguntas_respondidas(doc, numeros_respondidos):
    p = doc.add_paragraph(style="DDQ Corpo")
    p.paragraph_format.left_indent  = Cm(0.2)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(8)
    rotulo = p.add_run("Questions answered: ")
    rotulo.bold           = True
    rotulo.font.color.rgb = DDQ_COR_AZUL
    texto = ", ".join(numeros_respondidos) if numeros_respondidos else "none"
    r = p.add_run(texto)
    r.font.color.rgb = DDQ_COR_PRETO


# =============================================================================
# ── Etapa DDQ: excel → word (bytes) ─────────────────────────────────────────
# =============================================================================

def ddq_excel_para_word_bytes(
    excel_bytes: bytes,
    incluir_temas_vazios: bool = True,
    referencia_data: str = DDQ_REFERENCIA_DATA_DEFAULT,
) -> bytes:
    xls = pd.ExcelFile(io.BytesIO(excel_bytes), engine="openpyxl")
    doc = Document()
    ddq_configurar_documento(doc)

    ddq_adicionar_texto_vermelho(
        doc,
        (
            f"We were able to review the questions and will answer them as completely as "
            f"possible, by group, focusing on matters that have arisen or changed since "
            f"{referencia_data}. All values, unless otherwise stated, are in BRL. When "
            f"forecasts are mentioned, there is no assurance that such forecasts will "
            f"prevail, and it is likely that outcomes will vary from the forecasts."
        ),
        keep_with_next=True,
    )

    fallback_global = 1

    for sheet_name in xls.sheet_names:
        try:
            df = pd.read_excel(io.BytesIO(excel_bytes), sheet_name=sheet_name, engine="openpyxl")
        except Exception:
            if incluir_temas_vazios:
                p_tema = doc.add_paragraph(style="Heading 1")
                p_tema.paragraph_format.space_after = Pt(4)
                p_tema.add_run(sheet_name)
                ddq_adicionar_resumo_perguntas_respondidas(doc, [])
            continue

        if df.empty:
            if incluir_temas_vazios:
                p_tema = doc.add_paragraph(style="Heading 1")
                p_tema.paragraph_format.space_after = Pt(4)
                p_tema.add_run(sheet_name)
                ddq_adicionar_resumo_perguntas_respondidas(doc, [])
            continue

        try:
            cols = ddq_localizar_colunas(df)
        except Exception:
            if incluir_temas_vazios:
                p_tema = doc.add_paragraph(style="Heading 1")
                p_tema.paragraph_format.space_after = Pt(4)
                p_tema.add_run(sheet_name)
                ddq_adicionar_resumo_perguntas_respondidas(doc, [])
            continue

        df_filtrado = df[df[cols["include"]].apply(ddq_valor_incluir)].copy()

        if df_filtrado.empty and not incluir_temas_vazios:
            continue

        numeros_respondidos = ddq_montar_lista_perguntas_respondidas(
            df_filtrado, cols["number"]
        )

        p_tema = doc.add_paragraph(style="Heading 1")
        p_tema.paragraph_format.space_after = Pt(4)
        p_tema.add_run(sheet_name)

        ddq_adicionar_resumo_perguntas_respondidas(doc, numeros_respondidos)

        if numeros_respondidos:
            ddq_adicionar_texto_vermelho(
                doc,
                f"The remaining questions do not apply or no significant matters have "
                f"arisen or changed since {referencia_data}.",
            )
        else:
            ddq_adicionar_texto_vermelho(
                doc,
                f"No significant matters have arisen or changed since {referencia_data}.",
            )

        if df_filtrado.empty:
            continue

        for _, row in df_filtrado.iterrows():
            numero_real = ddq_normalizar_numero_pergunta(row.get(cols["number"]))
            pergunta    = ddq_limpar_texto(row.get(cols["question"]))
            resposta    = ddq_limpar_texto(row.get(cols["answer"]))

            ddq_adicionar_pergunta(
                doc,
                numero_real=numero_real,
                pergunta=pergunta,
                resposta=resposta,
                fallback_numero=fallback_global,
            )
            fallback_global += 1

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# =============================================================================
# ── UI ────────────────────────────────────────────────────────────────────────
# =============================================================================

tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs([
    "1 · .docx → .xlsx",
    "2 · Combinar excels",
    "3 · Aplicar fontes",
    "4 · .xlsx → .docx",
    "5 · Visualizar excel",
    "6 · DDQ → .docx",
])

# ── Tab 1 ──────────────────────────────────────────────────────────────────
with tab1:
    st.header("Etapa 1 – Extrair parágrafos do Word para Excel")
    st.write("Carrega um ficheiro **.docx** e gera um **.xlsx** com os parágrafos (col A) e as fontes/URLs detectadas (col B).")
    f1 = st.file_uploader("Ficheiro Word (.docx)", type=["docx"], key="t1_docx")
    if f1:
        with st.spinner("Extraindo parágrafos…"):
            out_bytes = word_to_excel_bytes(f1.read())
        st.success("Extração concluída!")
        st.download_button(
            "⬇️ Baixar Excel",
            data=out_bytes,
            file_name="RD_sourcebook_v0.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True,
        )

# ── Tab 2 ──────────────────────────────────────────────────────────────────
with tab2:
    st.header("Etapa 2 – Combinar ficheiros Excel")
    st.write("Empilha o Excel **novo** abaixo do Excel **antigo** (acumulado), gerando um único ficheiro.")
    col_n, col_a = st.columns(2)
    with col_n:
        f2_novo   = st.file_uploader("Excel NOVO (.xlsx)", type=["xlsx"], key="t2_novo")
    with col_a:
        f2_antigo = st.file_uploader("Excel ANTIGO / acumulado (.xlsx)", type=["xlsx"], key="t2_antigo")
    if st.button("▶️ Combinar", type="primary", use_container_width=True, key="t2_btn"):
        if not f2_novo or not f2_antigo:
            st.error("❌ Carregue os dois ficheiros.")
        else:
            with st.spinner("Combinando…"):
                combined = combinar_excels(f2_novo.read(), f2_antigo.read())
            st.success("Combinação concluída!")
            st.download_button(
                "⬇️ Baixar Excel combinado",
                data=combined,
                file_name="RD_sourcebook_matched_sources.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True,
                key="t2_dl",
            )

# ── Tab 3 ──────────────────────────────────────────────────────────────────
with tab3:
    st.header("Etapa 3 – Aplicar fontes no documento Word")
    st.write(
        "Carrega o Word alvo e o Excel acumulado com as fontes. "
        "O app usa **fuzzy matching** (rapidfuzz) para encontrar o parágrafo mais similar e "
        "insere a fonte em vermelho logo após."
    )

    col_w, col_x = st.columns(2)
    with col_w:
        f_word  = st.file_uploader("Ficheiro Word (.docx)", type=["docx"], key="t3_docx")
    with col_x:
        f_excel = st.file_uploader("Excel com fontes (.xlsx)", type=["xlsx"], key="t3_xlsx")

    threshold_val = st.slider(
        "Threshold de similaridade (fuzzy match)",
        min_value=50, max_value=100, value=80, step=5,
        key="t3_threshold",
        help="Score mínimo para aceitar um match. 100 = match quase exato.",
    )
    st.info(
        "ℹ️ **Threshold**: valores mais baixos → mais matches (menos precisos). "
        "Valores mais altos → menos matches (mais precisos). "
        "Score calculado com `partial_ratio` do rapidfuzz."
    )

    btn_t3 = st.button("▶️ Aplicar fontes", type="primary", use_container_width=True, key="t3_btn")

    if btn_t3:
        if not f_word or not f_excel:
            st.error("❌ Carregue o Word e o Excel antes de continuar.")
        else:
            with st.spinner("Fazendo matching e inserindo fontes…"):
                result_bytes, df_prev = apply_sources_bytes(
                    f_word.read(), f_excel.read(), threshold=threshold_val
                )
            st.session_state["t3_result"]  = result_bytes
            st.session_state["t3_preview"] = df_prev

    if "t3_result" in st.session_state:
        df_prev = st.session_state["t3_preview"]
        n_com   = int((df_prev["Fonte aplicada"] != "").sum())
        n_sem   = int((df_prev["Fonte aplicada"] == "").sum())
        m1, m2, m3 = st.columns(3)
        m1.metric("Parágrafos verificados", len(df_prev))
        m2.metric("Com fonte aplicada ✅", n_com)
        m3.metric("Sem match ❌", n_sem)

        st.markdown("#### 🔍 Resultado do matching")
        st.dataframe(
            df_prev,
            use_container_width=True,
            hide_index=True,
            column_config={
                "Parágrafo":      st.column_config.TextColumn("Parágrafo",      width="large"),
                "Fonte aplicada": st.column_config.TextColumn("Fonte aplicada", width="medium"),
                "Score":          st.column_config.NumberColumn("Score", format="%d", width="small"),
            },
        )

        st.download_button(
            "⬇️ Baixar Word com fontes",
            data=st.session_state["t3_result"],
            file_name="RD_v1.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            key="t3_dl",
        )

# ── Tab 4 ──────────────────────────────────────────────────────────────────
with tab4:
    st.header("Etapa 4 – Gerar Word com fontes a partir do Excel")
    st.write("Carrega o Excel acumulado (col A = texto, col B = fonte) e gera um **.docx** com texto e fontes em vermelho.")
    f4 = st.file_uploader("Excel (.xlsx)", type=["xlsx"], key="t4_xlsx")
    if st.button("▶️ Gerar Word", type="primary", use_container_width=True, key="t4_btn"):
        if not f4:
            st.error("❌ Carregue o Excel.")
        else:
            with st.spinner("Gerando Word…"):
                docx_out = excel_para_docx_bytes(f4.read())
            st.success("Word gerado!")
            st.download_button(
                "⬇️ Baixar Word",
                data=docx_out,
                file_name="RD_with_sources.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                key="t4_dl",
            )

# ── Tab 5 ──────────────────────────────────────────────────────────────────
with tab5:
    st.header("Visualizar Excel")
    f5 = st.file_uploader("Excel (.xlsx)", type=["xlsx"], key="t5_xlsx")
    if f5:
        df5 = pd.read_excel(io.BytesIO(f5.read()), header=None, engine="openpyxl")
        df5.columns = [f"Col {i+1}" for i in range(df5.shape[1])]
        st.dataframe(df5, use_container_width=True, hide_index=True)
        st.caption(f"{len(df5)} linhas · {df5.shape[1]} colunas")

# ── Tab 6 – DDQ ────────────────────────────────────────────────────────────
with tab6:
    st.header("DDQ – Gerar Word formatado a partir do Excel")
    st.write(
        "Carrega o ficheiro Excel do DDQ (**.xlsx**). "
        "Cada aba do Excel corresponde a um tema/secção. "
        "Apenas as linhas com **Include = 1** são incluídas no Word."
    )

    f6 = st.file_uploader(
        "Ficheiro Excel DDQ (.xlsx)",
        type=["xlsx"],
        key="t6_xlsx",
    )

    with st.expander("⚙️ Opções avançadas", expanded=False):
        col_ref, col_inc = st.columns(2)
        with col_ref:
            ref_data = st.text_input(
                "Data de referência",
                value=DDQ_REFERENCIA_DATA_DEFAULT,
                key="t6_refdata",
                help="Texto usado nas frases de contextualização (ex.: 'February 2026').",
            )
        with col_inc:
            incluir_vazios = st.checkbox(
                "Incluir temas sem perguntas respondidas",
                value=True,
                key="t6_vazios",
            )

    if st.button("▶️ Gerar Word DDQ", type="primary", use_container_width=True, key="t6_btn"):
        if not f6:
            st.error("❌ Carregue o ficheiro Excel DDQ.")
        else:
            with st.spinner("Gerando documento Word DDQ…"):
                try:
                    docx_ddq = ddq_excel_para_word_bytes(
                        excel_bytes=f6.read(),
                        incluir_temas_vazios=incluir_vazios,
                        referencia_data=ref_data,
                    )
                    st.session_state["t6_result"] = docx_ddq
                    st.success("Documento DDQ gerado com sucesso!")
                except Exception as e:
                    st.error(f"❌ Erro ao gerar o documento: {e}")

    if "t6_result" in st.session_state:
        st.download_button(
            "⬇️ Baixar Word DDQ",
            data=st.session_state["t6_result"],
            file_name=f"DDQ_{date.today().isoformat()}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            key="t6_dl",
        )
